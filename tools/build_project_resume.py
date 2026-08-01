from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_OUT = ROOT / "output" / "resume" / "Isaac_Hernandez_Project_Resume.docx"

INK = RGBColor(22, 34, 55)
ACCENT = RGBColor(48, 88, 160)
MUTED = RGBColor(79, 91, 108)
BLACK = RGBColor(25, 25, 25)
FONT = "Arial"


def set_run_font(run, size, *, bold=False, color=BLACK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def set_repeat_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def section_heading(doc, text):
    p = doc.add_paragraph(style="Resume Section")
    p.add_run(text.upper())
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="Resume Bullet")
    p.add_run(text)
    return p


def add_project(doc, name, descriptor, bullets):
    p = doc.add_paragraph(style="Resume Project")
    r = p.add_run(name)
    set_run_font(r, 9.8, bold=True, color=INK)
    r = p.add_run(f"  |  {descriptor}")
    set_run_font(r, 9.6, color=MUTED)
    for bullet in bullets:
        add_bullet(doc, bullet)


def build():
    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.66)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.8)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.4)
    normal.paragraph_format.line_spacing = 1.04

    section_style = styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
    section_style.base_style = normal
    section_style.font.name = FONT
    section_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    section_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    section_style.font.size = Pt(10.1)
    section_style.font.bold = True
    section_style.font.color.rgb = ACCENT
    section_style.paragraph_format.space_before = Pt(8)
    section_style.paragraph_format.space_after = Pt(3.5)
    section_style.paragraph_format.keep_with_next = True

    project_style = styles.add_style("Resume Project", WD_STYLE_TYPE.PARAGRAPH)
    project_style.base_style = normal
    project_style.paragraph_format.space_before = Pt(3.5)
    project_style.paragraph_format.space_after = Pt(1.5)
    project_style.paragraph_format.keep_with_next = True

    bullet_style = styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.base_style = styles["List Bullet"]
    bullet_style.font.name = FONT
    bullet_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    bullet_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    bullet_style.font.size = Pt(9.7)
    bullet_style.font.color.rgb = BLACK
    bullet_style.paragraph_format.left_indent = Inches(0.2)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.13)
    bullet_style.paragraph_format.space_before = Pt(0)
    bullet_style.paragraph_format.space_after = Pt(1.8)
    bullet_style.paragraph_format.line_spacing = 1.02

    # Plain-paragraph header: ATS and accessibility tools read the identity and
    # contact details in their intended order without interpreting a layout
    # table as tabular data.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("ISAAC HERNANDEZ")
    set_run_font(r, 23, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    r = p.add_run("FOUNDER & AI SYSTEMS ENGINEER")
    set_run_font(r, 9.7, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("isaacsight@gmail.com  |  github.com/isaacsight  |  kernel.chat")
    set_run_font(r, 9.2, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(
        "Founder and sole engineer behind an open-source AI agent ecosystem spanning "
        "multi-provider execution, MCP, provenance, agent security, orchestration, and "
        "a production React/Supabase platform. Builds reliable systems around models: "
        "deterministic computation, scoped authority, human approval, and replayable audit trails."
    )
    set_run_font(r, 10.0, color=BLACK)

    section_heading(doc, "Technical strengths")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("Core: ")
    set_run_font(r, 9.6, bold=True, color=INK)
    r = p.add_run(
        "TypeScript, JavaScript, React, Node.js, SQL, PostgreSQL, Supabase, REST APIs, Docker, GitHub Actions"
    )
    set_run_font(r, 9.6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("AI systems: ")
    set_run_font(r, 9.6, bold=True, color=INK)
    r = p.add_run(
        "agent orchestration, MCP, tool use, RAG, memory, evaluations, model routing, local inference, auditability"
    )
    set_run_font(r, 9.6)

    section_heading(doc, "Experience")
    p = doc.add_paragraph(style="Resume Project")
    r = p.add_run("KERNEL / KERNEL.CHAT")
    set_run_font(r, 10.3, bold=True, color=INK)
    r = p.add_run("  |  Founder & Sole Engineer")
    set_run_font(r, 9.7, bold=True, color=MUTED)
    r = p.add_run("  |  Nov 2025 - Present")
    set_run_font(r, 9.5, color=MUTED)

    bullets = [
        "Built and published kbot, an MIT-licensed terminal AI agent with 35 specialist agents, 20 model providers, 600+ tools, local/offline execution, an SDK, and an MCP server.",
        "Designed the runtime across routing, task planning, persistent memory, retrieval, tool approval, offline caching, scheduling, security, evaluation, and multi-agent orchestration.",
        "Shipped a React 19/TypeScript web platform backed by Supabase edge functions and PostgreSQL, with authentication, rate limiting, billing infrastructure, PWA support, internationalization, and Capacitor mobile packaging.",
        "Maintained a 12-package monorepo with Vitest and Playwright coverage, GitHub Actions CI, Docker support, automated npm publishing, technical RFCs, threat models, and customer-facing documentation.",
    ]
    for bullet in bullets:
        add_bullet(doc, bullet)

    section_heading(doc, "Selected open-source systems")
    add_project(
        doc,
        "@kernel.chat/kbot-finance",
        "Apache-2.0 provenance infrastructure",
        [
            "Separated probabilistic AI reasoning from deterministic source-of-truth computation using typed engine adapters and content-addressed request/response envelopes.",
            "Implemented hash-chained audit logs, integrity verification, replay, human approval gates, jurisdiction-aware verifier rules, reason codes, and an MCP server.",
        ],
    )
    add_project(
        doc,
        "@kernel.chat/agent-os",
        "OS primitives for AI agents",
        [
            "Built signed capabilities, per-agent token/cost/time quotas, taint-aware tool execution, downscoped handoff, credential isolation, namespaces, snapshots, and rubric-graded outcomes.",
        ],
    )
    add_project(
        doc,
        "@kernel.chat/kbot-orchestrator",
        "Auditable agent-human workflows",
        [
            "Created resumable pipelines with explicit approval gates, typed state transitions, audit trails, and reusable outreach, content, and maintenance workflows.",
        ],
    )

    section_heading(doc, "Public work")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "Published packages, source, architecture notes, RFCs, security analyses, and the kernel.chat technical magazine at "
    )
    set_run_font(r, 9.6)
    r = p.add_run("github.com/isaacsight/kernel")
    set_run_font(r, 9.6, bold=True, color=ACCENT)
    r = p.add_run(". Work is directly runnable through npm and documented for external contributors and adopters.")
    set_run_font(r, 9.6)

    # Quiet footer; no page-number field is needed for a one-page resume.
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("Open-source portfolio: github.com/isaacsight/kernel")
    set_run_font(r, 8, color=MUTED)

    core = doc.core_properties
    core.title = "Isaac Hernandez - Project Resume"
    core.subject = "AI systems engineering and open-source agent infrastructure"
    core.author = "Isaac Hernandez"
    core.keywords = "AI agents, TypeScript, MCP, provenance, audit, open source"
    core.comments = "Generated from verified public project information."

    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    build()
